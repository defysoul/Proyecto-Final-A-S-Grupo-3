"""Skill de documentos (Paso 8): genera CATASTRO y EVIDENCIA en .docx utilizando plantillas institucionales."""

from __future__ import annotations

import asyncio
import base64
import binascii
import difflib
import os
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from ..client import SisavClient
from .core import obtener_detalle_iniciativa, listar_postulaciones


def _normalizar_columna(texto: str) -> str:
    """Normaliza el título de una columna (colapsa saltos de línea/espacios y pasa a minúsculas) para poder
    identificarla sin importar cómo esté formateado exactamente en la plantilla."""
    return " ".join((texto or "").replace("\n", " ").split()).strip().casefold()


# Mapea el nombre normalizado de cada columna del CATASTRO (tal como aparece en la fila de
# encabezado de la plantilla) a la función que extrae su valor desde la postulación. Las
# columnas que no aparecen aquí (p. ej. "Área", "Unidad") quedan en blanco para completarse
# manualmente.
_CATASTRO_CAMPOS_CONOCIDOS: dict[str, Any] = {
    "año": lambda fila: str(fila.get("anio") or ""),
    "código sisav": lambda fila: str(fila.get("idpostulacion", "")),
    "proyecto": lambda fila: fila.get("nombre") or "",
    "tipo de proyecto": lambda fila: (fila.get("convocatoria") or {}).get("nombre") or "",
}

_EVIDENCIA_ASISTENCIA_COLUMNAS = ["Nombre", "Cédula de Identidad", "Correo"]

_LIMIT_PAGINA = 100


def _docs_dir() -> Path:
    """Carpeta local donde se escriben los .docx generados (no se sube a SISAV2)."""
    home = Path.home()
    for candidato in ("Desktop", "Escritorio", "Documents", "Documentos"):
        destino = home / candidato / "SISAV2 - Documentos generados"
        if (home / candidato).exists():
            destino.mkdir(parents=True, exist_ok=True)
            return destino
    destino = home / "SISAV2 - Documentos generados"
    destino.mkdir(parents=True, exist_ok=True)
    return destino


def _slug(texto: str) -> str:
    limpio = "".join(c if c.isalnum() else "_" for c in texto.strip())
    while "__" in limpio:
        limpio = limpio.replace("__", "_")
    return limpio.strip("_") or "sin_nombre"


def _leer_encabezados_catastro(tabla: Any) -> list[str]:
    """Lee el nombre de cada columna directamente desde la fila de encabezado (fila 2 de la
    plantilla, índice 1) en vez de asumir un listado fijo."""
    if len(tabla.rows) < 2:
        return []
    return [celda.text for celda in tabla.rows[1].cells]


def _rellenar_fila_catastro(celdas: Any, fila: dict[str, Any], columnas: list[str]) -> None:
    """Escribe los valores de una postulación en una fila de la tabla, ubicando cada dato según
    el nombre real de la columna (leído de la plantilla) en vez de una posición fija."""
    for idx, celda in enumerate(celdas):
        nombre_columna = _normalizar_columna(columnas[idx]) if idx < len(columnas) else ""
        obtener_valor = _CATASTRO_CAMPOS_CONOCIDOS.get(nombre_columna)
        _set_cell_text(celda, obtener_valor(fila) if obtener_valor else "")


def _mejor_coincidencia_carrera(
    nombre_buscado: str, candidatos: set[str]
) -> tuple[str | None, list[str]]:
    """Encuentra, entre los nombres de carrera que aparecen en los datos, el más parecido al que se pidió."""
    buscado = nombre_buscado.strip().casefold()
    mapa = {c.strip().casefold(): c for c in candidatos if c and c.strip()}
    if buscado in mapa:
        return mapa[buscado], []
    cercanos = difflib.get_close_matches(buscado, list(mapa.keys()), n=3, cutoff=0.6)
    if cercanos:
        return mapa[cercanos[0]], [mapa[c] for c in cercanos]
    return None, []


def _set_cell_text(cell: Any, texto: str, *, bold: bool = False, size: int | None = None) -> None:
    """Inserta texto conservando estrictamente el formato XML (fuente, tamaño, color) del run original."""
    rPr_clonado = None
    
    # Extraer las propiedades XML (Run Properties) del molde antes de limpiarlo
    if cell.paragraphs and cell.paragraphs[0].runs:
        r0 = cell.paragraphs[0].runs[0]
        if r0._element.rPr is not None:
            rPr_clonado = deepcopy(r0._element.rPr)
            
    # Limpiar todos los párrafos de la celda
    for p in cell.paragraphs:
        p.clear()
        
    # Remover párrafos adicionales para evitar saltos de línea indeseados
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
        
    p0 = cell.paragraphs[0]
    run = p0.add_run(texto)
    
    # Restaurar el formato original inyectando el nodo XML guardado
    if rPr_clonado is not None:
        run._element.append(rPr_clonado)
        
    if bold:
        run.bold = True
    if size is not None:
        run.font.size = Pt(size)


# --------------------------------------------------------------------------- #
# Descripción automática (a partir del form-builder de la postulación)
# --------------------------------------------------------------------------- #

_ETIQUETAS_OBJETIVO_DESCRIPCION = (
    "objetivo",
    "descripción",
    "descripcion",
    "necesidad, problema u oportunidad",
)


def _campo(obj: Any, nombre: str, default: Any = None) -> Any:
    if obj is None:
        return default
    if isinstance(obj, dict):
        return obj.get(nombre, default)
    return getattr(obj, nombre, default)


def _extraer_descripcion_por_defecto(detalle: Any) -> str | None:
    """Arma una descripción por defecto a partir de los campos tipo textarea de la postulación."""
    pasos = _campo(detalle, "pasos", []) or []
    coincidencias: list[str] = []
    respaldo: str | None = None

    for paso in pasos:
        campos = _campo(paso, "campos", []) or []
        for campo in campos:
            if (_campo(campo, "type") or "").strip().lower() != "textarea":
                continue
            valor = _campo(campo, "value")
            if not valor or not str(valor).strip():
                continue
            etiqueta = (_campo(campo, "label") or "").strip().lower()
            if respaldo is None:
                respaldo = str(valor).strip()
            if any(clave in etiqueta for clave in _ETIQUETAS_OBJETIVO_DESCRIPCION):
                coincidencias.append(str(valor).strip())

    if coincidencias:
        return "\n\n".join(coincidencias)
    return respaldo


# --------------------------------------------------------------------------- #
# CATASTRO
# --------------------------------------------------------------------------- #


async def generar_catastro(
    client: SisavClient,
    *,
    id_usuario: int,
    roles: list[int],
    carrera_id: int,
    carrera_nombre: str,
    facultad_nombre: str,
    anio: int | None = None,
    modalidad: str = "PRE_GRADO",
) -> dict[str, Any]:
    """Arma el CATASTRO de una carrera en .docx utilizando plantilla_catastro.docx con retardo y reintentos."""
    
    ruta_plantilla = Path(__file__).parent / "plantilla_catastro.docx"
    if not ruta_plantilla.exists():
        return {
            "error_critico": f"Archivo de plantilla no encontrado en: {ruta_plantilla}. "
                             "Asegúrese de que plantilla_catastro.docx esté exactamente "
                             "en la misma carpeta que documents.py."
        }
        
    filas_anio: list[dict[str, Any]] = []
    offset = 0
    
    while True:
        intentos = 3
        pagina = None
        
        while intentos > 0:
            try:
                pagina = await listar_postulaciones(
                    client,
                    modalidad=modalidad,
                    id_usuario=id_usuario,
                    roles=roles,
                    carrera=carrera_id,
                    ingreso=True,
                    offset=offset,
                    limit=_LIMIT_PAGINA,
                )
                break
            except Exception as e:
                intentos -= 1
                if intentos == 0:
                    raise e
                await asyncio.sleep(1)
        
        for p in pagina["postulaciones"]:
            if anio is None or p.get("anio") == anio:
                filas_anio.append(p)
                
        offset += _LIMIT_PAGINA
        if offset >= pagina["total"]:
            break
            
        await asyncio.sleep(0.5)

    candidatos = {p.get("carrera", "") for p in filas_anio}
    carrera_resuelta, alternativas = _mejor_coincidencia_carrera(carrera_nombre, candidatos)

    if carrera_resuelta is None:
        filas: list[dict[str, Any]] = []
    else:
        objetivo = carrera_resuelta.strip().casefold()
        filas = [
            p for p in filas_anio
            if (p.get("carrera") or "").strip().casefold() == objetivo
        ]

    filas.sort(key=lambda f: (int(f.get("anio") or 0), int(f.get("idpostulacion") or 0)))

    if anio is not None:
        anio_mostrar = str(anio)
    else:
        anio_maximo = max((int(f.get("anio") or 0) for f in filas), default=0)
        anio_mostrar = str(anio_maximo) if anio_maximo > 0 else "Sin registros"

    nombre_para_documento = carrera_resuelta or carrera_nombre

    documento = Document(str(ruta_plantilla))

    # Reemplazo de metadatos de forma global (párrafos superiores y tablas)
    reemplazos = {
        "{{ANIO}}": anio_mostrar,
        "{{CARRERA}}": nombre_para_documento,
        "{{FACULTAD}}": facultad_nombre,
    }

    for parrafo in documento.paragraphs:
        for run in parrafo.runs:
            for clave, valor in reemplazos.items():
                if clave in run.text:
                    run.text = run.text.replace(clave, valor)

    for tabla_doc in documento.tables:
        for fila_doc in tabla_doc.rows:
            for celda_doc in fila_doc.cells:
                for parrafo in celda_doc.paragraphs:
                    for run in parrafo.runs:
                        for clave, valor in reemplazos.items():
                            if clave in run.text:
                                run.text = run.text.replace(clave, valor)

    tabla = documento.tables[0]

    # Fila 2 de la plantilla (índice 1): encabezado real, se lee tal cual está en el archivo y
    # no se toca ni se usa como molde de datos.
    columnas = _leer_encabezados_catastro(tabla)

    # Fila 3 en adelante (índice 2+): contenido de ejemplo. Se limpia el texto manteniendo el
    # formato de celda y se conserva solo la primera como molde para clonar; si hubiera más de
    # una fila de ejemplo, las sobrantes se eliminan.
    filas_ejemplo = list(tabla.rows[2:])
    fila_plantilla = None
    if filas_ejemplo:
        fila_plantilla = filas_ejemplo[0]
        for celda in fila_plantilla.cells:
            _set_cell_text(celda, "")
        for fila_extra in filas_ejemplo[1:]:
            tabla._tbl.remove(fila_extra._tr)

    # Inserción de filas de datos clonando la fila molde (o agregando filas nuevas si la
    # plantilla no traía ninguna fila de ejemplo bajo el encabezado).
    if fila_plantilla is not None:
        tr_plantilla = fila_plantilla._tr

        if not filas:
            tabla._tbl.remove(tr_plantilla)
        else:
            for i, fila in enumerate(filas):
                if i == 0:
                    celdas = fila_plantilla.cells
                else:
                    nuevo_tr = deepcopy(tr_plantilla)
                    tabla._tbl.append(nuevo_tr)
                    celdas = tabla.rows[-1].cells

                _rellenar_fila_catastro(celdas, fila, columnas)
    else:
        for fila in filas:
            celdas = tabla.add_row().cells
            _rellenar_fila_catastro(celdas, fila, columnas)

    nombre_archivo = f"CATASTRO_{_slug(nombre_para_documento)}_{anio_mostrar}.docx"
    ruta = _docs_dir() / nombre_archivo
    documento.save(ruta)

    campos_pendientes = [
        nombre_columna
        for nombre_columna in columnas
        if _normalizar_columna(nombre_columna) not in _CATASTRO_CAMPOS_CONOCIDOS
    ]

    resultado: dict[str, Any] = {
        "archivo": str(ruta),
        "carrera_solicitada": carrera_nombre,
        "carrera_detectada": carrera_resuelta,
        "anio_reportado": anio_mostrar,
        "iniciativas_incluidas": len(filas),
        "campos_pendientes_de_completar_manualmente": campos_pendientes,
    }
    if carrera_resuelta is None:
        resultado["advertencia"] = (
            "No se encontró ninguna carrera parecida a "
            f"'{carrera_nombre}' en las iniciativas. El documento se generó vacío."
        )
    elif carrera_resuelta.strip().casefold() != carrera_nombre.strip().casefold():
        resultado["advertencia"] = (
            f"Se pidió '{carrera_nombre}' pero se usó '{carrera_resuelta}', "
            "que es el nombre más parecido encontrado en los datos reales."
        )
    if alternativas:
        resultado["otras_carreras_similares"] = alternativas
    return resultado


# --------------------------------------------------------------------------- #
# EVIDENCIA
# --------------------------------------------------------------------------- #


async def generar_evidencia(
    client: SisavClient,
    *,
    id_postulacion: int,
    nombre_analista: str,
    descripcion: str | None = None,
    fotos_base64: list[str] | None = None,
    asistentes: list[dict[str, str]] | None = None,
    difusion: str | None = None,
    resumen_asistencia: str | None = None,
) -> dict[str, Any]:
    """Arma el informe de EVIDENCIA leyendo variables desde tablas y párrafos en la plantilla."""
    
    ruta_plantilla = Path(__file__).parent / "plantilla_evidencia.docx"
    if not ruta_plantilla.exists():
        return {
            "error_critico": f"Archivo de plantilla no encontrado en: {ruta_plantilla}. "
                             "Asegúrese de que plantilla_evidencia.docx esté exactamente "
                             "en la misma carpeta que documents.py."
        }
        
    documento = Document(str(ruta_plantilla))
    detalle = await obtener_detalle_iniciativa(client, id_postulacion=id_postulacion)

    descripcion_generada_automaticamente = False
    if not descripcion or not descripcion.strip():
        descripcion = _extraer_descripcion_por_defecto(detalle)
        descripcion_generada_automaticamente = descripcion is not None
        if not descripcion:
            descripcion = "Sin descripción disponible. Complete este apartado manualmente."

    nombre_iniciativa = _campo(detalle, "nombre") or f"Iniciativa {id_postulacion}"
    fecha_iniciativa = _campo(detalle, "fecha") or ""
    facultad_iniciativa = _campo(_campo(detalle, "facultad", {}), "nombre") or "No informada"
    carrera_iniciativa = _campo(_campo(detalle, "carrera", {}), "nombre") or "No informada"

    reemplazos = {
        "{{CODIGO}}": str(id_postulacion),
        "{{FECHA}}": str(fecha_iniciativa).split("T")[0],
        "{{NOMBRE_INICIATIVA}}": nombre_iniciativa,
        "{{FACULTAD}}": facultad_iniciativa,
        "{{CARRERA}}": carrera_iniciativa,
        "{{ANALISTA}}": nombre_analista,
    }

    for parrafo in documento.paragraphs:
        for run in parrafo.runs:
            for clave, valor in reemplazos.items():
                if clave in run.text:
                    run.text = run.text.replace(clave, valor)

    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        for clave, valor in reemplazos.items():
                            if clave in run.text:
                                run.text = run.text.replace(clave, valor)

    documento.add_heading("DESCRIPCIÓN DE LA INICIATIVA", level=2)
    documento.add_paragraph(descripcion)

    documento.add_heading("DIFUSIÓN", level=2)
    documento.add_paragraph(difusion or "Esta iniciativa no requiere difusión.")

    documento.add_heading("FOTOGRAFÍAS DE LA INICIATIVA", level=2)
    fotos_incluidas = 0
    fotos_con_error = 0
    for i, foto_b64 in enumerate(fotos_base64 or []):
        try:
            datos = base64.b64decode(foto_b64, validate=True)
        except (binascii.Error, ValueError):
            fotos_con_error += 1
            continue
        tmp = _docs_dir() / f"_tmp_foto_{id_postulacion}_{i}.png"
        tmp.write_bytes(datos)
        try:
            documento.add_picture(str(tmp), width=Cm(14))
        finally:
            tmp.unlink(missing_ok=True)
        fotos_incluidas += 1

    documento.add_heading("LISTA DE ASISTENCIA", level=2)
    asistentes = asistentes or []
    tabla_asistencia = documento.add_table(rows=1, cols=len(_EVIDENCIA_ASISTENCIA_COLUMNAS))
    tabla_asistencia.style = "Table Grid"
    for celda, titulo_col in zip(tabla_asistencia.rows[0].cells, _EVIDENCIA_ASISTENCIA_COLUMNAS):
        _set_cell_text(celda, titulo_col, bold=True)
        
    for persona in asistentes:
        celdas = tabla_asistencia.add_row().cells
        _set_cell_text(celdas[0], persona.get("nombre", ""))
        _set_cell_text(celdas[1], persona.get("cedula", ""))
        _set_cell_text(celdas[2], persona.get("correo", ""))
        
    if resumen_asistencia:
        documento.add_paragraph(resumen_asistencia)

    documento.add_heading("Validación de Medios de Verificación", level=2)
    documento.add_paragraph(
        "En mi calidad de Vicerrectora de Vinculación con el Medio, de la Universidad "
        "Tecnológica Metropolitana, he revisado los medios de verificación asociados a las "
        "actividades expuestas en el marco del proyecto en curso. Tras el análisis "
        "correspondiente, se valida que dichos medios cumplen con los criterios establecidos "
        "para evidenciar la ejecución y cumplimiento de las acciones comprometidas. Esta "
        "validación se emite para efectos de control interno y seguimiento del proyecto, y se "
        "pone a disposición de los organismos pertinentes para su revisión y respaldo "
        "documental."
    )
    documento.add_paragraph("__________________________\nVicerrectora de Vinculación con el Medio")

    nombre_archivo = f"EVIDENCIA_{id_postulacion}_{_slug(nombre_iniciativa)}.docx"
    ruta = _docs_dir() / nombre_archivo
    documento.save(ruta)

    resultado: dict[str, Any] = {
        "archivo": str(ruta),
        "iniciativa": nombre_iniciativa,
        "codigo_sisav": id_postulacion,
        "fotos_incluidas": fotos_incluidas,
        "fotos_con_error": fotos_con_error,
        "asistentes_incluidos": len(asistentes),
        "descripcion_generada_automaticamente": descripcion_generada_automaticamente,
    }
    if descripcion_generada_automaticamente:
        resultado["advertencia_descripcion"] = (
            "La descripción se extrajo del formulario de postulación. Verifique "
            "su contenido antes del reporte definitivo."
        )
    if not asistentes:
        resultado["advertencia_asistencia"] = (
            "No se entregaron asistentes mediante los parámetros de la herramienta."
        )
    return resultado